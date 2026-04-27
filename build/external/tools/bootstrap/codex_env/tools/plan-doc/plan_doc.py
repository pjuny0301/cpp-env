#!/usr/bin/env python3
from __future__ import annotations

import argparse
import base64
import hashlib
import json
import os
import re
import sys
import tempfile
import textwrap
import urllib.error
import urllib.request
from pathlib import Path
from typing import Any

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

VERSION = "0.2.1"
DEFAULT_MODEL = "gpt-image-1.5"


def is_windows() -> bool:
    return os.name == "nt"


def print_output(payload: dict[str, Any], as_json: bool) -> None:
    if as_json:
        print(json.dumps(payload, indent=2, ensure_ascii=False))
        return

    status = payload.get("status", "ok")
    print(f"status: {status}")
    for key, value in payload.items():
        if key == "status":
            continue
        if isinstance(value, (dict, list)):
            print(f"{key}: {json.dumps(value, indent=2, ensure_ascii=False)}")
        else:
            print(f"{key}: {value}")


def default_plan_template_root() -> Path:
    if "PLAN_TEMPLATE_ROOT" in os.environ and os.environ["PLAN_TEMPLATE_ROOT"]:
        return Path(os.environ["PLAN_TEMPLATE_ROOT"])
    return Path("C:/dev/templates") if is_windows() else Path.home() / "dev" / "templates"


def default_theme() -> dict[str, Any]:
    return {
        "document": {
            "font_name": "Malgun Gothic" if is_windows() else "DejaVu Sans",
            "accent_hex": "1F4E79",
            "secondary_hex": "16324F",
            "muted_hex": "617B94",
        },
        "image_style": {
            "cover": "cinematic technical editorial art, polished planning deck visual, clean whitespace, layered lighting",
            "section": "bold editorial systems illustration, presentation-ready technical concept art, clear focal point",
            "inline": "infographic-inspired concept illustration, concrete software architecture motifs, restrained but polished",
        },
    }


def load_theme() -> dict[str, Any]:
    theme_path = default_plan_template_root() / "plan-doc" / "theme.json"
    if theme_path.exists():
        return json.loads(theme_path.read_text(encoding="utf-8"))
    return default_theme()


def ensure_plan_root(project_root: Path) -> Path:
    plan_root = project_root / "plan"
    plan_root.mkdir(parents=True, exist_ok=True)
    return plan_root


def bundle_root_for_source(source_path: Path) -> Path:
    return source_path.parent / source_path.stem


def read_json_file(path: Path) -> dict[str, Any] | None:
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as error:
        raise RuntimeError(f"Invalid JSON in {path}: {error}") from error


def resolve_source(project_root: Path, source: str | None) -> Path:
    if source:
        path = Path(source)
        if not path.is_absolute():
            path = project_root / source
        if path.exists():
            return path
        raise SystemExit(f"Plan source not found: {path}")

    for candidate in ("plan/plan.md", "plan/plan.txt"):
        path = project_root / candidate
        if path.exists():
            return path
    raise SystemExit("Could not find plan/plan.md or plan/plan.txt.")


def parse_markdown_plan(text: str) -> dict[str, Any]:
    lines = text.splitlines()
    title = ""
    sections: list[dict[str, Any]] = []
    current_title: str | None = None
    current_lines: list[str] = []

    def flush_section() -> None:
        nonlocal current_title, current_lines
        if not current_title:
            return
        paragraphs = normalize_paragraphs(current_lines)
        sections.append({"title": current_title, "paragraphs": paragraphs})
        current_title = None
        current_lines = []

    for raw_line in lines:
        line = raw_line.rstrip()
        if line.startswith("# "):
            if not title:
                title = line[2:].strip()
            else:
                flush_section()
                current_title = line[2:].strip()
        elif line.startswith("## "):
            flush_section()
            current_title = line[3:].strip()
        else:
            current_lines.append(line)

    flush_section()

    if not title:
        non_empty = [line.strip() for line in lines if line.strip()]
        title = non_empty[0] if non_empty else "Project Plan"

    if not sections:
        sections = [{"title": "Summary", "paragraphs": normalize_paragraphs(lines)}]

    return {"title": title, "sections": sections}


def normalize_paragraphs(lines: list[str]) -> list[str]:
    paragraphs: list[str] = []
    buffer: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if buffer:
                paragraphs.append(" ".join(buffer))
                buffer = []
            continue
        if stripped.startswith("- "):
            if buffer:
                paragraphs.append(" ".join(buffer))
                buffer = []
            paragraphs.append(stripped)
        else:
            buffer.append(stripped)
    if buffer:
        paragraphs.append(" ".join(buffer))
    return [paragraph for paragraph in paragraphs if paragraph]


def infer_role(title: str) -> str:
    lowered = title.lower()
    if "summary" in lowered or "개요" in title:
        return "summary"
    if "key" in lowered or "change" in lowered or "주요" in title:
        return "key_changes"
    if "arch" in lowered or "pipeline" in lowered or "구조" in title or "파이프라인" in title:
        return "architecture"
    if "test" in lowered or "검증" in title or "테스트" in title:
        return "testing"
    if "assumption" in lowered or "가정" in title:
        return "assumptions"
    return "section"


def slugify(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9]+", "-", value).strip("-").lower()
    return cleaned or "section"


def prompt_hash(prompt: str) -> str:
    return hashlib.sha256(prompt.encode("utf-8")).hexdigest()[:10]


def visual_filename(visual: dict[str, Any]) -> str:
    if visual["kind"] == "cover":
        return f"cover-hero-{prompt_hash(visual['prompt'])}.png"
    if visual["kind"] == "banner":
        return f"section-{slugify(visual['title'])}-{prompt_hash(visual['prompt'])}.png"
    return (
        f"inline-{slugify(visual['title'])}-{visual['paragraph_index'] + 1:02d}-"
        f"{prompt_hash(visual['prompt'])}.png"
    )


def layout_variant(role: str) -> str:
    return {
        "summary": "hero-summary",
        "key_changes": "stacked-highlights",
        "architecture": "system-diagram",
        "testing": "checklist-grid",
        "assumptions": "callout-panel",
    }.get(role, "section-banner")


def build_prompt(theme: dict[str, Any], title: str, excerpt: str, role: str, kind: str) -> str:
    style = theme["image_style"]["cover" if kind == "cover" else "section" if kind == "banner" else "inline"]
    sanitized_excerpt = excerpt[:320].replace("\n", " ").strip()
    return (
        f"{style}. Focus on '{title}'. "
        f"Support the planning role '{role}'. "
        f"Use visual metaphors grounded in the following plan text: {sanitized_excerpt}. "
        "High clarity, presentation-ready, avoid extra text labels unless they materially help understanding."
    )


def select_inline_paragraphs(sections: list[dict[str, Any]], max_inline_images: int) -> list[tuple[int, int]]:
    chosen: list[tuple[int, int]] = []
    for section_index, section in enumerate(sections):
        for paragraph_index, paragraph in enumerate(section["paragraphs"]):
            if len(chosen) >= max_inline_images:
                return chosen
            if len(paragraph) < 60:
                continue
            chosen.append((section_index, paragraph_index))
            break
    return chosen


def build_render_spec(parsed: dict[str, Any], theme: dict[str, Any], max_inline_images: int) -> dict[str, Any]:
    title = parsed["title"]
    sections = parsed["sections"]
    inline_targets = set(select_inline_paragraphs(sections, max_inline_images))

    cover_prompt = build_prompt(
        theme,
        title,
        " ".join(sections[0]["paragraphs"][:2]) if sections else title,
        "cover",
        "cover",
    )
    visuals: list[dict[str, Any]] = [
        {
            "id": "cover-hero",
            "kind": "cover",
            "title": title,
            "role": "cover",
            "layout_variant": "cover-hero",
            "filename": f"cover-hero-{prompt_hash(cover_prompt)}.png",
            "prompt": cover_prompt,
            "section_index": None,
            "paragraph_index": None,
        }
    ]

    spec_sections: list[dict[str, Any]] = []
    for index, section in enumerate(sections):
        role = infer_role(section["title"])
        banner_filename = f"section-{index + 1:02d}-banner.png"
        banner_prompt = build_prompt(
            theme,
            section["title"],
            " ".join(section["paragraphs"][:2]),
            role,
            "banner",
        )
        banner_filename = f"section-{slugify(section['title'])}-{prompt_hash(banner_prompt)}.png"
        banner_visual = {
            "id": f"section-{index + 1:02d}-banner",
            "kind": "banner",
            "title": section["title"],
            "role": role,
            "layout_variant": layout_variant(role),
            "filename": banner_filename,
            "prompt": banner_prompt,
            "section_index": index,
            "paragraph_index": None,
        }
        visuals.append(banner_visual)

        inline_images: list[dict[str, Any]] = []
        for paragraph_index, paragraph in enumerate(section["paragraphs"]):
            if (index, paragraph_index) not in inline_targets:
                continue
            prompt = build_prompt(theme, section["title"], paragraph, role, "inline")
            filename = (
                f"inline-{slugify(section['title'])}-{paragraph_index + 1:02d}-"
                f"{prompt_hash(prompt)}.png"
            )
            visual = {
                "id": f"section-{index + 1:02d}-inline-{paragraph_index + 1:02d}",
                "kind": "inline",
                "title": section["title"],
                "role": role,
                "layout_variant": "inline-concept",
                "filename": filename,
                "prompt": prompt,
                "section_index": index,
                "paragraph_index": paragraph_index,
            }
            visuals.append(visual)
            inline_images.append(visual)

        spec_sections.append(
            {
                "id": f"section-{index + 1:02d}",
                "title": section["title"],
                "role": role,
                "layout_variant": layout_variant(role),
                "paragraphs": section["paragraphs"],
                "banner": banner_visual,
                "inline_images": inline_images,
            }
        )

    return {
        "title": title,
        "theme": theme,
        "visuals": visuals,
        "sections": spec_sections,
    }


def prompt_override_path(bundle_root: Path) -> Path:
    return bundle_root / "prompt_overrides.json"


def bundle_manifest_path(bundle_root: Path) -> Path:
    return bundle_root / "bundle_manifest.json"


def placeholder_notice_path(bundle_root: Path) -> Path:
    return bundle_root / "placeholder_notice.txt"


def override_entries_from_payload(payload: dict[str, Any] | None) -> dict[str, dict[str, Any]]:
    if not payload:
        return {}

    raw_entries = payload.get("overrides")
    if raw_entries is None:
        raw_entries = payload.get("prompts")

    entries: dict[str, dict[str, Any]] = {}
    if isinstance(raw_entries, list):
        for entry in raw_entries:
            if not isinstance(entry, dict):
                continue
            entry_id = str(entry.get("id", "")).strip()
            if entry_id:
                entries[entry_id] = entry
    elif isinstance(raw_entries, dict):
        for entry_id, entry in raw_entries.items():
            if isinstance(entry, dict):
                entries[str(entry_id)] = entry
    return entries


def build_override_payload(bundle_root: Path, source_path: Path, spec: dict[str, Any]) -> dict[str, Any]:
    override_path = prompt_override_path(bundle_root)
    existing_payload = read_json_file(override_path)
    if existing_payload is None:
        existing_payload = read_json_file(bundle_root / "prompts.json")
    existing_entries = override_entries_from_payload(existing_payload)

    overrides: list[dict[str, Any]] = []
    for visual in spec["visuals"]:
        existing = existing_entries.get(visual["id"], {})
        prompt = str(existing.get("prompt") or visual["prompt"])
        layout = str(existing.get("layout_variant") or visual["layout_variant"])
        pinned_asset = str(existing.get("pinned_asset") or "").strip()
        note = str(existing.get("note") or "").strip()

        visual["prompt"] = prompt
        visual["layout_variant"] = layout
        if pinned_asset:
            visual["pinned_asset"] = pinned_asset
        else:
            visual.pop("pinned_asset", None)
        visual["filename"] = visual_filename(visual)

        overrides.append(
            {
                "id": visual["id"],
                "kind": visual["kind"],
                "title": visual["title"],
                "role": visual["role"],
                "layout_variant": layout,
                "prompt": prompt,
                "pinned_asset": pinned_asset,
                "note": note,
                "default_asset_bundle_relpath": f"assets/{visual['filename']}",
            }
        )

    return {
        "schema_version": 1,
        "source": str(source_path),
        "bundle_root": str(bundle_root),
        "instructions": [
            "Edit prompt to regenerate only that visual on the next run.",
            "Set pinned_asset to a file inside this bundle to reuse it without regeneration.",
            "Keep pinned_asset relative and inside this bundle so plan assets stay isolated from other plan files.",
        ],
        "overrides": overrides,
    }


def resolve_pinned_asset(bundle_root: Path, pinned_asset: str | None) -> Path | None:
    if not pinned_asset:
        return None
    pinned_path = Path(pinned_asset)
    if pinned_path.is_absolute():
        raise RuntimeError("pinned_asset must be relative to the bundle root.")

    candidate = (bundle_root / pinned_path).resolve()
    bundle_resolved = bundle_root.resolve()
    try:
        candidate.relative_to(bundle_resolved)
    except ValueError as error:
        raise RuntimeError("pinned_asset must stay inside the source-specific bundle root.") from error

    if not candidate.exists():
        return None
    return candidate


def choose_font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    candidates = []
    if is_windows():
        candidates.extend(
            [
                "C:/Windows/Fonts/malgunbd.ttf" if bold else "C:/Windows/Fonts/malgun.ttf",
                "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
            ]
        )
    candidates.extend(
        [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        ]
    )
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def hex_to_rgb(value: str) -> tuple[int, int, int]:
    value = value.strip().lstrip("#")
    return tuple(int(value[index:index + 2], 16) for index in (0, 2, 4))


def gradient_seed(prompt: str, accent_hex: str, secondary_hex: str) -> tuple[tuple[int, int, int], tuple[int, int, int]]:
    accent = list(hex_to_rgb(accent_hex))
    secondary = list(hex_to_rgb(secondary_hex))
    digest = hashlib.sha256(prompt.encode("utf-8")).digest()
    for index in range(3):
        accent[index] = min(255, max(0, accent[index] + digest[index] % 40 - 20))
        secondary[index] = min(255, max(0, secondary[index] + digest[index + 3] % 40 - 20))
    return tuple(accent), tuple(secondary)


def wrap_text(text: str, width: int) -> str:
    return "\n".join(textwrap.wrap(text, width=width, break_long_words=False))


def generate_placeholder_image(output_path: Path, visual: dict[str, Any], theme: dict[str, Any]) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    width = 1536 if visual["kind"] == "cover" else 1365
    height = 1024 if visual["kind"] == "cover" else 768

    accent, secondary = gradient_seed(
        visual["prompt"],
        theme["document"]["accent_hex"],
        theme["document"]["secondary_hex"],
    )

    image = Image.new("RGB", (width, height), secondary)
    draw = ImageDraw.Draw(image)
    for y in range(height):
        blend = y / max(1, height - 1)
        color = tuple(int(accent[index] * (1 - blend) + secondary[index] * blend) for index in range(3))
        draw.line((0, y, width, y), fill=color)

    overlay_color = tuple(min(255, channel + 28) for channel in secondary)
    draw.rounded_rectangle(
        (72, 72, width - 72, height - 72),
        radius=36,
        fill=overlay_color,
        outline=(255, 255, 255),
        width=3,
    )

    title_font = choose_font(54 if visual["kind"] == "cover" else 42, bold=True)
    body_font = choose_font(28)
    label_font = choose_font(24, bold=True)

    draw.text((120, 120), visual["kind"].upper(), fill=(255, 255, 255), font=label_font)
    draw.text((120, 190), wrap_text(visual["title"], 22), fill=(255, 255, 255), font=title_font, spacing=8)

    excerpt = visual["prompt"]
    excerpt = re.sub(r"\s+", " ", excerpt)
    excerpt = wrap_text(excerpt[:320], 38)
    draw.multiline_text((120, 330), excerpt, fill=(232, 240, 248), font=body_font, spacing=10)

    accent_bar = tuple(min(255, channel + 50) for channel in accent)
    draw.rounded_rectangle((120, height - 180, width - 120, height - 120), radius=22, fill=accent_bar)
    draw.text((148, height - 166), "Generated locally because no image API credentials were configured.", fill=(20, 27, 36), font=body_font)

    image.save(output_path)


def generate_openai_image(prompt: str, output_path: Path, model: str) -> None:
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY is not set.")

    payload = json.dumps(
        {
            "model": model,
            "prompt": prompt,
            "size": "1536x1024",
            "quality": "medium",
            "output_format": "png",
            "background": "auto",
        }
    ).encode("utf-8")

    request = urllib.request.Request(
        "https://api.openai.com/v1/images/generations",
        data=payload,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )

    try:
        with urllib.request.urlopen(request, timeout=180) as response:
            body = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as error:
        detail = error.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"OpenAI image request failed: {detail}") from error

    image_data = body.get("data", [])
    if not image_data or "b64_json" not in image_data[0]:
        raise RuntimeError("OpenAI image response did not contain base64 image data.")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_bytes(base64.b64decode(image_data[0]["b64_json"]))


def create_assets(
    bundle_root: Path,
    spec: dict[str, Any],
    provider: str,
    model: str,
) -> tuple[list[dict[str, Any]], str, list[dict[str, Any]], list[str]]:
    assets_root = bundle_root / "assets"
    assets_root.mkdir(parents=True, exist_ok=True)

    chosen_provider = provider
    provider_notes: list[str] = []
    if provider == "auto":
        if os.environ.get("OPENAI_API_KEY"):
            chosen_provider = "openai"
        else:
            chosen_provider = "placeholder"
            provider_notes.append(
                "OPENAI_API_KEY was not set, so plan-doc generated placeholder graphics automatically."
            )

    results: list[dict[str, Any]] = []
    placeholder_assets: list[dict[str, Any]] = []
    for visual in spec["visuals"]:
        pinned_asset = resolve_pinned_asset(bundle_root, visual.get("pinned_asset"))
        if pinned_asset is not None:
            output_path = pinned_asset
            used_provider = "pinned"
        else:
            output_path = assets_root / visual["filename"]
            used_provider = chosen_provider

        if pinned_asset is None and output_path.exists():
            used_provider = "reused"
        elif pinned_asset is None and chosen_provider == "openai":
            try:
                generate_openai_image(visual["prompt"], output_path, model)
            except Exception:
                generate_placeholder_image(output_path, visual, spec["theme"])
                used_provider = "placeholder"
                provider_notes.append(
                    f"OpenAI image generation failed for {visual['id']}; a placeholder graphic was generated instead."
                )
        elif pinned_asset is None:
            generate_placeholder_image(output_path, visual, spec["theme"])
            used_provider = "placeholder"

        asset_record = {
            **visual,
            "asset_path": str(output_path),
            "asset_relpath": str(output_path.relative_to(bundle_root.parent)),
            "asset_bundle_relpath": str(output_path.relative_to(bundle_root)),
            "provider": used_provider,
            "pinned_asset": visual.get("pinned_asset") or None,
        }
        results.append(asset_record)
        if used_provider == "placeholder":
            placeholder_assets.append(asset_record)

    deduped_notes: list[str] = []
    for note in provider_notes:
        if note not in deduped_notes:
            deduped_notes.append(note)

    return results, chosen_provider, placeholder_assets, deduped_notes


def apply_doc_style(document: Document, theme: dict[str, Any]) -> None:
    font_name = theme["document"]["font_name"]
    normal_style = document.styles["Normal"]
    normal_style.font.name = font_name
    normal_style.font.size = Pt(10.5)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    for style_name, size in (("Title", 24), ("Heading 1", 18), ("Heading 2", 16)):
        style = document.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def add_paragraph(document: Document, text: str, theme: dict[str, Any], bullet: bool = False) -> None:
    paragraph = document.add_paragraph(style="Normal")
    if bullet:
        paragraph.style = document.styles["List Bullet"]
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text[2:].strip() if bullet and text.startswith("- ") else text)
    run.font.name = theme["document"]["font_name"]
    run._element.rPr.rFonts.set(qn("w:eastAsia"), theme["document"]["font_name"])
    run.font.size = Pt(10.5)


def compose_docx(
    output_path: Path,
    title: str,
    sections: list[dict[str, Any]],
    assets: list[dict[str, Any]],
    theme: dict[str, Any],
    source_path: Path,
) -> None:
    asset_lookup = {asset["id"]: asset for asset in assets}
    accent = RGBColor(*hex_to_rgb(theme["document"]["accent_hex"]))
    muted = RGBColor(*hex_to_rgb(theme["document"]["muted_hex"]))

    document = Document()
    apply_doc_style(document, theme)

    cover = asset_lookup["cover-hero"]
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = accent
    title_run.font.name = theme["document"]["font_name"]
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), theme["document"]["font_name"])

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"Generated from {source_path.name}")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = muted
    subtitle_run.font.name = theme["document"]["font_name"]
    subtitle_run._element.rPr.rFonts.set(qn("w:eastAsia"), theme["document"]["font_name"])

    document.add_picture(cover["asset_path"], width=Inches(6.6))
    cover_caption = document.add_paragraph()
    cover_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_run = cover_caption.add_run("Cover concept")
    cover_run.italic = True
    cover_run.font.color.rgb = muted
    cover_run.font.name = theme["document"]["font_name"]
    cover_run._element.rPr.rFonts.set(qn("w:eastAsia"), theme["document"]["font_name"])

    document.add_section(WD_SECTION.NEW_PAGE)

    for section in sections:
        heading = document.add_heading(section["title"], level=1)
        heading.runs[0].font.color.rgb = accent
        heading.runs[0].font.name = theme["document"]["font_name"]
        heading.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), theme["document"]["font_name"])

        banner = asset_lookup[section["banner"]["id"]]
        document.add_picture(banner["asset_path"], width=Inches(6.4))

        for paragraph_index, paragraph_text in enumerate(section["paragraphs"]):
            add_paragraph(document, paragraph_text, theme, bullet=paragraph_text.startswith("- "))

            inline_match = next(
                (
                    image
                    for image in section["inline_images"]
                    if image["paragraph_index"] == paragraph_index
                ),
                None,
            )
            if inline_match:
                inline_asset = asset_lookup[inline_match["id"]]
                document.add_picture(inline_asset["asset_path"], width=Inches(5.8))
                caption = document.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = caption.add_run(section["title"])
                run.italic = True
                run.font.color.rgb = muted
                run.font.name = theme["document"]["font_name"]
                run._element.rPr.rFonts.set(qn("w:eastAsia"), theme["document"]["font_name"])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def generate_plan_doc(
    project_root: Path,
    source: str | None,
    provider: str,
    image_model: str,
    max_inline_images: int,
) -> dict[str, Any]:
    project_root = project_root.resolve()
    plan_root = ensure_plan_root(project_root)
    source_path = resolve_source(project_root, source)
    bundle_root = bundle_root_for_source(source_path)
    bundle_root.mkdir(parents=True, exist_ok=True)
    theme = load_theme()
    parsed = parse_markdown_plan(source_path.read_text(encoding="utf-8"))
    spec = build_render_spec(parsed, theme, max_inline_images)
    override_payload = build_override_payload(bundle_root, source_path, spec)
    assets, provider_used, placeholder_assets, provider_notes = create_assets(
        bundle_root,
        spec,
        provider,
        image_model,
    )
    asset_lookup = {asset["id"]: asset for asset in assets}

    for entry in override_payload["overrides"]:
        asset = asset_lookup.get(entry["id"])
        if not asset:
            continue
        entry["asset_bundle_relpath"] = asset["asset_bundle_relpath"]
        entry["provider"] = asset["provider"]

    render_spec_payload = {
        "title": spec["title"],
        "source": str(source_path),
        "bundle_root": str(bundle_root),
        "prompt_overrides": str(prompt_override_path(bundle_root)),
        "sections": spec["sections"],
        "assets": assets,
    }
    prompts_payload = {
        "provider_requested": provider,
        "provider_used": provider_used,
        "image_model": image_model,
        "prompt_override_file": str(prompt_override_path(bundle_root)),
        "prompts": [
            {
                "id": asset["id"],
                "kind": asset["kind"],
                "title": asset["title"],
                "prompt": asset["prompt"],
                "asset_relpath": asset["asset_relpath"],
                "asset_bundle_relpath": asset["asset_bundle_relpath"],
                "layout_variant": asset["layout_variant"],
                "pinned_asset": asset["pinned_asset"],
            }
            for asset in assets
        ],
    }

    bundle_manifest_payload = {
        "schema_version": 1,
        "source": str(source_path),
        "bundle_root": str(bundle_root),
        "editable_files": {
            "source": str(source_path),
            "prompt_overrides": str(prompt_override_path(bundle_root)),
            "render_spec": str(bundle_root / "render_spec.json"),
            "prompts": str(bundle_root / "prompts.json"),
            "docx": str(bundle_root / "plan_summary.docx"),
        },
        "asset_policy": {
            "scope": "source-local bundle",
            "reusable": True,
            "isolation": "assets stay inside the source-specific bundle and should not be shared across sibling bundles",
            "default_reuse_key": "prompt hash",
            "pinned_asset_behavior": "reuse an existing file inside the same bundle without regeneration",
        },
        "placeholder_generation": {
            "automatic_placeholder_mode": bool(provider_notes),
            "detected_this_run": bool(placeholder_assets),
            "placeholder_count": len(placeholder_assets),
            "notice_file": str(placeholder_notice_path(bundle_root)) if (placeholder_assets or provider_notes) else None,
            "assets": [
                {
                    "id": asset["id"],
                    "title": asset["title"],
                    "asset_bundle_relpath": asset["asset_bundle_relpath"],
                }
                for asset in placeholder_assets
            ],
            "reasons": provider_notes,
        },
        "token_budget_guidance": {
            "read_first": ["bundle_manifest.json", "prompt_overrides.json"],
            "read_second": ["render_spec.json", "prompts.json"],
            "avoid_until_needed": ["plan_summary.docx"],
        },
        "assets": [
            {
                "id": asset["id"],
                "kind": asset["kind"],
                "title": asset["title"],
                "asset_bundle_relpath": asset["asset_bundle_relpath"],
                "provider": asset["provider"],
            }
            for asset in assets
        ],
    }

    render_spec_path = bundle_root / "render_spec.json"
    prompts_path = bundle_root / "prompts.json"
    output_docx = bundle_root / "plan_summary.docx"
    source_snapshot = bundle_root / f"source_snapshot{source_path.suffix}"
    override_path = prompt_override_path(bundle_root)
    manifest_path = bundle_manifest_path(bundle_root)
    notice_path = placeholder_notice_path(bundle_root)

    render_spec_path.write_text(json.dumps(render_spec_payload, indent=2, ensure_ascii=False), encoding="utf-8")
    prompts_path.write_text(json.dumps(prompts_payload, indent=2, ensure_ascii=False), encoding="utf-8")
    override_path.write_text(json.dumps(override_payload, indent=2, ensure_ascii=False), encoding="utf-8")
    manifest_path.write_text(json.dumps(bundle_manifest_payload, indent=2, ensure_ascii=False), encoding="utf-8")
    source_snapshot.write_text(source_path.read_text(encoding="utf-8"), encoding="utf-8")
    if placeholder_assets or provider_notes:
        notice_lines = [
            "Automatic placeholder fallback information for this plan bundle.",
            "",
            f"Source: {source_path}",
            f"Bundle: {bundle_root}",
            f"Placeholder assets generated this run: {len(placeholder_assets)}",
        ]
        if provider_notes:
            notice_lines.extend(["", "Reasons:"])
            notice_lines.extend(f"- {note}" for note in provider_notes)
        if placeholder_assets:
            notice_lines.extend(["", "Assets:"])
            notice_lines.extend(
                f"- {asset['id']}: {asset['asset_bundle_relpath']}" for asset in placeholder_assets
            )
        notice_lines.extend(
            [
                "",
                "Re-run plan-doc with a working OpenAI image configuration to replace placeholder outputs when needed.",
            ]
        )
        notice_path.write_text("\n".join(notice_lines) + "\n", encoding="utf-8")
    elif notice_path.exists():
        notice_path.unlink()
    compose_docx(output_docx, parsed["title"], spec["sections"], assets, theme, source_path)

    return {
        "status": "ok",
        "project_root": str(project_root),
        "source": str(source_path),
        "bundle_root": str(bundle_root),
        "provider_used": provider_used,
        "render_spec": str(render_spec_path),
        "prompts": str(prompts_path),
        "prompt_overrides": str(override_path),
        "bundle_manifest": str(manifest_path),
        "docx": str(output_docx),
        "asset_count": len(assets),
        "placeholder_generated": bool(placeholder_assets),
        "placeholder_mode": bool(provider_notes),
        "placeholder_count": len(placeholder_assets),
        "placeholder_notice": str(notice_path) if (placeholder_assets or provider_notes) else None,
        "warnings": provider_notes,
    }


def self_test() -> dict[str, Any]:
    sample_plan = textwrap.dedent(
        """\
        # Sample Plan

        ## Summary
        - Establish a Windows-first but Linux-portable project standard.
        - Generate a polished plan document from markdown.

        ## Architecture
        The source of truth should live in CMake, vcpkg, and platform-separated source folders.

        ## Test Plan
        Verify tool output, scaffold generation, and runtime deployment.
        """
    )

    with tempfile.TemporaryDirectory(prefix="plan-doc-") as temp_dir:
        project_root = Path(temp_dir)
        plan_root = project_root / "plan"
        plan_root.mkdir(parents=True, exist_ok=True)
        source_path = plan_root / "plan.md"
        source_path.write_text(sample_plan, encoding="utf-8")

        result = generate_plan_doc(
            project_root=project_root,
            source=str(source_path),
            provider="placeholder",
            image_model=DEFAULT_MODEL,
            max_inline_images=2,
        )

        required = [
            Path(result["render_spec"]),
            Path(result["prompts"]),
            Path(result["docx"]),
        ]
        missing = [str(path) for path in required if not path.exists()]
        if missing:
            raise SystemExit(f"Self-test failed. Missing outputs: {missing}")

        Document(result["docx"])
        return result


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Generate a polished plan_summary.docx and related assets from project plan text."
    )
    parser.add_argument("project_root", nargs="?", default=".", help="Project root containing plan/plan.md or plan/plan.txt.")
    parser.add_argument("--source", help="Explicit source path relative to project root or absolute path.")
    parser.add_argument(
        "--image-provider",
        choices=("auto", "openai", "placeholder"),
        default="auto",
        help="Image generation provider.",
    )
    parser.add_argument(
        "--image-model",
        default=os.environ.get("OPENAI_IMAGE_MODEL", DEFAULT_MODEL),
        help="OpenAI image model when --image-provider=openai or auto.",
    )
    parser.add_argument("--max-inline-images", type=int, default=4, help="Maximum inline concept images to generate.")
    parser.add_argument("--self-test", action="store_true", help="Run tool self-test.")
    parser.add_argument("--version", action="store_true", help="Print tool version.")
    parser.add_argument("--json", action="store_true", help="Emit machine-readable JSON.")
    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()

    if args.version:
        print(VERSION)
        return 0

    try:
        if args.self_test:
            print_output(self_test(), args.json)
            return 0

        result = generate_plan_doc(
            project_root=Path(args.project_root),
            source=args.source,
            provider=args.image_provider,
            image_model=args.image_model,
            max_inline_images=args.max_inline_images,
        )
        print_output(result, args.json)
        return 0
    except SystemExit:
        raise
    except Exception as error:
        print_output({"status": "error", "message": str(error)}, args.json)
        return 1


if __name__ == "__main__":
    sys.exit(main())
